#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Risk Filter for Shortlisted Candidates
- Detects hidden text in resumes (white-on-white, hidden runs)
- Identifies job hopping (3+ jobs in the last 2 years)
- Identifies employment gaps > 2 years
- Removes risky candidates from Shortlisted.xlsx
- Saves a separate file with risky candidates and their risk reasons
- Uses Candidate Name to locate resume file in the Resumes folder
- Fixed date parsing to support both abbreviated and full month names
"""

import os
import re
import logging
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from datetime import datetime, date
from dotenv import load_dotenv

load_dotenv()

# ============================================
# CONFIGURATION FROM ENVIRONMENT
# ============================================
SHORTLISTED_FILE = os.getenv("SHORTLISTED_FILE", "Shortlisted.xlsx")
RESUMES_FOLDER = os.getenv("RESUMES_FOLDER", "Resumes")
OUTPUT_FILE = os.getenv("RISK_FILTERED_FILE", "Shortlisted_clean.xlsx")
RISKY_FILE = os.getenv("RISKY_FILE", "Risky_Candidates.xlsx")

LOG_FILE = "risk_filter.log"

# Risk thresholds from environment (with defaults)
MAX_JOBS_IN_LAST_YEARS = int(os.getenv("MAX_JOBS_IN_LAST_YEARS", 2))
MAX_JOB_HOP_COUNT = int(os.getenv("MAX_JOB_HOP_COUNT", 3))
MAX_GAP_MONTHS = int(os.getenv("MAX_GAP_MONTHS", 24))

# ============================================
# LOGGING
# ============================================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE, encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# ============================================
# HELPER FUNCTIONS
# ============================================

def find_resume_file(candidate_name):
    """
    Search the RESUMES_FOLDER for a .docx file whose name contains the candidate's name
    (case-insensitive). Returns the full file path if found, else None.
    """
    if not os.path.isdir(RESUMES_FOLDER):
        logger.error(f"Resumes folder not found: {RESUMES_FOLDER}")
        return None

    name_parts = candidate_name.lower().split()
    for file in os.listdir(RESUMES_FOLDER):
        if not file.lower().endswith('.docx'):
            continue
        base = os.path.splitext(file)[0].lower()
        if all(part in base for part in name_parts):
            return os.path.join(RESUMES_FOLDER, file)
    for file in os.listdir(RESUMES_FOLDER):
        if not file.lower().endswith('.docx'):
            continue
        if candidate_name.lower() in file.lower():
            return os.path.join(RESUMES_FOLDER, file)
    return None

def extract_text_with_formatting(docx_path):
    """
    Extract paragraphs and runs with formatting info.
    Returns list of (text, font_color_rgb, is_hidden).
    """
    try:
        doc = Document(docx_path)
        runs_info = []
        for para in doc.paragraphs:
            for run in para.runs:
                text = run.text.strip()
                if not text:
                    continue
                is_hidden = run.font.hidden if run.font.hidden is not None else False
                color = None
                if run.font.color and run.font.color.rgb:
                    color = run.font.color.rgb
                runs_info.append((text, color, is_hidden))
        return runs_info
    except Exception as e:
        logger.error(f"Error reading {docx_path}: {e}")
        return []

def detect_hidden_text(runs_info):
    """
    Returns True if any hidden text or white-on-white found.
    """
    white = RGBColor(255, 255, 255)
    for text, color, hidden in runs_info:
        if hidden:
            logger.info(f"  Hidden text detected (property hidden): {text[:50]}...")
            return True
        if color == white:
            logger.info(f"  White-on-white text detected: {text[:50]}...")
            return True
    return False

def parse_dates_from_text(text):
    """
    Extracts date ranges like "Jan 2020 – Mar 2022" or "June 2017 – December 2019".
    Returns list of (start_date, end_date) tuples as date objects.
    Now supports both abbreviated (e.g., "Jan") and full (e.g., "January") month names.
    """
    patterns = [
        r'([A-Za-z]{3,9}\s+\d{4})\s*[-–—]\s*([A-Za-z]{3,9}\s+\d{4}|Present)',
        r'(\d{4}[-/]\d{2})\s*[-–—]\s*(\d{4}[-/]\d{2}|Present)',
        r'(\d{4})\s*[-–—]\s*(\d{4}|Present)'
    ]
    dates = []
    for pat in patterns:
        matches = re.findall(pat, text, re.IGNORECASE)
        for match in matches:
            start_str, end_str = match[0], match[1]
            try:
                # Parse start date
                if re.match(r'[A-Za-z]{3,9}\s+\d{4}', start_str):
                    # Try abbreviated month first, then full month
                    try:
                        start = datetime.strptime(start_str, '%b %Y').date()
                    except ValueError:
                        try:
                            start = datetime.strptime(start_str, '%B %Y').date()
                        except ValueError:
                            continue
                elif re.match(r'\d{4}[-/]\d{2}', start_str):
                    start = datetime.strptime(start_str, '%Y-%m').date()
                elif re.match(r'\d{4}', start_str):
                    start = datetime.strptime(start_str + '-01-01', '%Y-%m-%d').date()
                else:
                    continue

                # Parse end date
                if end_str.lower() == 'present':
                    end = date.today()
                elif re.match(r'[A-Za-z]{3,9}\s+\d{4}', end_str):
                    try:
                        end = datetime.strptime(end_str, '%b %Y').date()
                    except ValueError:
                        try:
                            end = datetime.strptime(end_str, '%B %Y').date()
                        except ValueError:
                            continue
                elif re.match(r'\d{4}[-/]\d{2}', end_str):
                    end = datetime.strptime(end_str, '%Y-%m').date()
                elif re.match(r'\d{4}', end_str):
                    end = datetime.strptime(end_str + '-12-31', '%Y-%m-%d').date()
                else:
                    continue

                dates.append((start, end))
            except Exception as e:
                logger.debug(f"Date parsing failed: {e}")
                continue
    return dates

def analyze_employment_history(text):
    """
    Returns (job_hop_risk, gap_risk) booleans.
    """
    dates = parse_dates_from_text(text)
    if len(dates) < 2:
        return False, False

    dates.sort(key=lambda x: x[0])

    cutoff = date.today().replace(year=date.today().year - MAX_JOBS_IN_LAST_YEARS)
    recent_jobs = [d for d in dates if d[0] >= cutoff]
    job_hop_risk = len(recent_jobs) > MAX_JOB_HOP_COUNT
    if job_hop_risk:
        logger.info(f"  Job hopping risk: {len(recent_jobs)} jobs in last {MAX_JOBS_IN_LAST_YEARS} years.")

    gap_risk = False
    for i in range(1, len(dates)):
        prev_end = dates[i-1][1]
        curr_start = dates[i][0]
        gap_months = (curr_start - prev_end).days / 30.44
        if gap_months > MAX_GAP_MONTHS:
            logger.info(f"  Gap > {MAX_GAP_MONTHS} months detected: {gap_months:.1f} months between jobs.")
            gap_risk = True
            break

    return job_hop_risk, gap_risk

def process_resume(filepath):
    """
    Run all risk checks on a resume.
    Returns dict: {'hidden': bool, 'job_hop': bool, 'gap': bool}
    """
    risks = {'hidden': False, 'job_hop': False, 'gap': False}

    runs_info = extract_text_with_formatting(filepath)
    risks['hidden'] = detect_hidden_text(runs_info)

    try:
        doc = Document(filepath)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
    except:
        full_text = ""

    job_hop, gap = analyze_employment_history(full_text)
    risks['job_hop'] = job_hop
    risks['gap'] = gap

    return risks

# ============================================
# MAIN FILTERING LOGIC
# ============================================
def main():
    logger.info("="*60)
    logger.info("RISK FILTER STARTED")
    logger.info("="*60)

    if not os.path.exists(SHORTLISTED_FILE):
        logger.error(f"Shortlisted file not found: {SHORTLISTED_FILE}")
        return

    df_short = pd.read_excel(SHORTLISTED_FILE, engine='openpyxl')

    name_col = None
    for col in df_short.columns:
        if 'name' in col.lower():
            name_col = col
            break
    if name_col is None:
        logger.error("No column containing 'name' found in shortlisted file.")
        return

    logger.info(f"Loaded {len(df_short)} shortlisted candidates. Using name column: '{name_col}'")

    risky_indices = []
    risk_details = []
    risky_rows = []

    for idx, row in df_short.iterrows():
        candidate_name = row[name_col]
        filepath = find_resume_file(candidate_name)
        if filepath is None:
            logger.warning(f"No resume file found for candidate: {candidate_name}. Skipping risk check.")
            continue

        logger.info(f"\nChecking {candidate_name} ({os.path.basename(filepath)})...")
        risks = process_resume(filepath)

        if any(risks.values()):
            risky_indices.append(idx)
            risk_summary = []
            if risks['hidden']: risk_summary.append('hidden text')
            if risks['job_hop']: risk_summary.append('job hopping')
            if risks['gap']: risk_summary.append('gap >2y')
            risk_str = ', '.join(risk_summary)
            risk_details.append(f"{candidate_name}: {risk_str}")
            logger.info(f"  - RISK DETECTED: {risk_str}")

            risky_row = row.to_dict()
            risky_row['Risk Factors'] = risk_str
            risky_rows.append(risky_row)
        else:
            logger.info("  - No risks detected.")

    if risky_rows:
        df_risky = pd.DataFrame(risky_rows)
        original_cols = list(df_short.columns)
        df_risky = df_risky[original_cols + ['Risk Factors']]
        df_risky.to_excel(RISKY_FILE, index=False)
        logger.info(f"\nRisky candidates saved to {RISKY_FILE} ({len(df_risky)} candidates)")
    else:
        logger.info("\nNo risky candidates found.")

    if risky_indices:
        logger.info(f"\nRemoving {len(risky_indices)} risky candidates from shortlist.")
        df_clean = df_short.drop(index=risky_indices).reset_index(drop=True)
        df_clean.to_excel(OUTPUT_FILE, index=False)
        logger.info(f"SUCCESS: Clean shortlist saved to {OUTPUT_FILE} ({len(df_clean)} candidates)")
    else:
        logger.info("\nSUCCESS: No risky candidates found. Shortlist unchanged.")
        df_short.to_excel(OUTPUT_FILE, index=False)
        logger.info(f"Shortlist copied to {OUTPUT_FILE}")

    logger.info("="*60)
    logger.info("RISK FILTER COMPLETED")
    logger.info("="*60)

if __name__ == "__main__":
    main()