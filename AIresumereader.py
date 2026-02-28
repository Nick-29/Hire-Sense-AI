#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
AI-Powered Resume Screening & Ranking System
- Reads job requisitions from Excel.
- Processes .docx resumes from a local folder.
- Extracts candidate name, email, phone, experience, and text.
- Uses hybrid matching (exact phrase + semantic similarity).
- Computes weighted score (skills + experience) and ranks candidates per job.
- Outputs Shortlisted.xlsx and Non_Shortlisted.xlsx.
"""

import os
import re
import zipfile
import pandas as pd
import numpy as np
from docx import Document
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv

load_dotenv()

# ==============================
# CONFIGURATION FROM ENVIRONMENT
# ==============================
BASE_PATH = os.getenv("BASE_PATH", os.path.dirname(__file__))
SKILLS_FILE = os.getenv("SKILLS_FILE", os.path.join(BASE_PATH, "Open Positions Skill Set Details.xlsx"))
RESUMES_FOLDER = os.getenv("RESUMES_FOLDER", os.path.join(BASE_PATH, "Resumes"))
SHORTLISTED_FILE = os.getenv("SHORTLISTED_FILE", os.path.join(BASE_PATH, "Shortlisted.xlsx"))
NON_SHORTLISTED_FILE = os.getenv("NON_SHORTLISTED_FILE", os.path.join(BASE_PATH, "Non_Shortlisted.xlsx"))

# Scoring weights (can also be set in .env)
SKILL_WEIGHT = float(os.getenv("SKILL_WEIGHT", 0.7))
EXPERIENCE_WEIGHT = float(os.getenv("EXPERIENCE_WEIGHT", 0.3))
SEMANTIC_THRESHOLD = float(os.getenv("SEMANTIC_THRESHOLD", 0.3))
EXACT_MATCH_SCORE = float(os.getenv("EXACT_MATCH_SCORE", 0.95))
SHORTLIST_THRESHOLD = float(os.getenv("SHORTLIST_THRESHOLD", 0.6))

# ==============================
# CREATE FOLDERS IF NOT EXIST
# ==============================
os.makedirs(RESUMES_FOLDER, exist_ok=True)

# ==============================
# LOAD AI MODEL
# ==============================
print("\n Loading AI model (all-mpnet-base-v2)...")
model = SentenceTransformer('all-mpnet-base-v2')
print("Model loaded.\n")

# ==============================
# HELPER FUNCTIONS
# ==============================
def extract_docx_text_deep(docx_path):
    """Extract text from .docx including paragraphs, tables, and hidden XML."""
    full_text = []
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                full_text.append(" ".join(cell.text for cell in row.cells))
        with zipfile.ZipFile(docx_path, 'r') as z:
            for xml_file in [f for f in z.namelist() if f.endswith('.xml')]:
                with z.open(xml_file) as f:
                    content = f.read().decode('utf-8', errors='ignore')
                    full_text.append(re.sub(r'<[^>]+>', ' ', content))
    except Exception as e:
        print(f"Error extracting {docx_path}: {e}")
    return "\n".join(full_text)

def extract_name(text, filename):
    """Extract candidate name from resume content."""
    if not text:
        return os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ').title()
    lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 2]
    indicators = ['name:', 'candidate name:', 'candidate:', 'name of candidate:']
    for line in lines[:15]:
        lower = line.lower()
        for ind in indicators:
            if ind in lower:
                name_part = line.split(':', 1)[1].strip() if ':' in line else line
                name_part = re.sub(r'[^\w\s]', '', name_part)
                words = name_part.split()
                if 1 <= len(words) <= 4 and not any(c.isdigit() for c in name_part):
                    return name_part.title()
    for line in lines[:10]:
        clean = re.sub(r'[#*_\-]', '', line).strip()
        words = clean.split()
        if 2 <= len(words) <= 4 and not any(c.isdigit() for c in clean):
            lower = line.lower()
            skip = ['resume', 'curriculum', 'vitae', 'cv', 'application',
                    'profile', 'summary', 'contact', 'email', 'phone',
                    'address', 'objective', 'education', 'experience',
                    'skills', 'projects', 'certifications', 'languages']
            if not any(s in lower for s in skip):
                if all(w[0].isupper() if w else False for w in words):
                    return clean.title()
    for line in lines[:10]:
        clean = re.sub(r'[#*_\-]', '', line).strip()
        if clean.isupper() and len(clean.split()) in [2, 3]:
            if not any(c.isdigit() for c in clean):
                return clean.title()
    return os.path.splitext(filename)[0].replace('_', ' ').replace('-', ' ').title()

def extract_phone_robust(text):
    """Extract 10-digit Indian phone number."""
    if not text:
        return "Not Found"
    text_no_dates = re.sub(r'\b\d{1,4}[-/\.]\d{1,2}[-/\.]\d{2,4}\b', ' ', text)
    patterns = [
        r'\+91[-\s]*([6-9]\d{9})\b',
        r'91[-\s]*([6-9]\d{9})\b',
        r'0[-\s]*([6-9]\d{9})\b',
        r'\b([6-9]\d{9})\b'
    ]
    for pat in patterns:
        match = re.search(pat, text_no_dates)
        if match:
            return match.group(1)
    return "Not Found"

def extract_email(text):
    emails = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text.lower())
    return emails[0] if emails else "Not Found"

def extract_experience(text):
    patterns = [
        r'(\d+\.?\d*)\s*\+?\s*(?:yrs?|years?)\s*(?:of)?\s*experience',
        r'experience\s*(?:of)?\s*(\d+\.?\d*)\s*\+?\s*(?:yrs?|years?)',
        r'over\s*(\d+\.?\d*)\s*\+?\s*(?:yrs?|years?)',
        r'(\d+\.?\d*)\s*\+?\s*(?:yrs?|years?)\s*(?:in|with)'
    ]
    text_lower = text.lower()
    for pat in patterns:
        match = re.search(pat, text_lower)
        if match:
            return float(match.group(1))
    return 0.0

def parse_years(exp_str):
    if pd.isna(exp_str) or not exp_str:
        return 0.0
    match = re.search(r'(\d+\.?\d*)', str(exp_str))
    return float(match.group(1)) if match else 0.0

def extract_skills_section(text):
    """Extract everything under 'SKILLS' heading until next major heading."""
    lines = text.split('\n')
    skills = []
    in_skills = False
    for line in lines:
        if re.search(r'^skills?\s*$', line, re.IGNORECASE):
            in_skills = True
            continue
        if in_skills:
            if re.search(r'^(education|experience|projects|certifications|technical skills|summary|objective)', line, re.IGNORECASE):
                break
            if line.strip():
                skills.append(line.strip())
    return " ".join(skills)

def hybrid_skill_match(resume_text, required_skills, skills_section, debug=True):
    """Hybrid exact + semantic skill matching."""
    if not required_skills:
        return [], 0.0
    if isinstance(required_skills, str):
        required_skills = [s.strip() for s in re.split(r'[,\n]', required_skills) if s.strip()]

    combined = resume_text[:500] + " " + skills_section
    resume_emb = model.encode(combined)
    skill_embs = model.encode(required_skills)
    similarities = cosine_similarity([resume_emb], skill_embs)[0]

    exact_matches = []
    for i, skill in enumerate(required_skills):
        escaped = re.escape(skill)
        pattern = r'\b' + escaped + r'\b'
        if re.search(pattern, resume_text, re.IGNORECASE):
            similarities[i] = EXACT_MATCH_SCORE
            exact_matches.append(skill)

    matched = []
    if debug:
        print("\n   Skill details (after exact-match check):")
    for skill, sim in zip(required_skills, similarities):
        status = "MATCHED" if sim >= SEMANTIC_THRESHOLD else "NO MATCH"
        if debug:
            print(f"      [{status}] {skill}: {sim:.3f}")
        if sim >= SEMANTIC_THRESHOLD:
            matched.append((skill, sim))

    avg_score = np.mean([sim for _, sim in matched]) if matched else 0.0
    if debug:
        print(f"   MATCHED {len(matched)} skills, average score: {avg_score:.3f}")
        if exact_matches:
            print(f"      (Exact matches: {', '.join(exact_matches)})")
    return matched, avg_score

def experience_score(candidate_years, required_min):
    if required_min <= 0:
        return 1.0
    if candidate_years >= required_min:
        return 1.0
    return candidate_years / required_min

def compute_overall_score(skill_score, exp_score):
    return skill_score * SKILL_WEIGHT + exp_score * EXPERIENCE_WEIGHT

# ==============================
# MAIN PROCESSING
# ==============================
def main():
    if not os.path.exists(SKILLS_FILE):
        print(f"ERROR: Skills file not found: {SKILLS_FILE}")
        return

    print("Loading requisitions...")
    try:
        df_req = pd.read_excel(SKILLS_FILE, engine='openpyxl')
        requisition_col = next((c for c in df_req.columns if 'requisition' in c.lower()), None)
        skill_col = next((c for c in df_req.columns if 'skill' in c.lower()), None)
        experience_col = next((c for c in df_req.columns if 'experience' in c.lower()), None)

        if not requisition_col or not skill_col:
            print("ERROR: Required columns not found.")
            return

        print(f"   Using columns: '{requisition_col}', '{skill_col}', '{experience_col if experience_col else 'None'}'")

        job_skills = {}
        for _, row in df_req.iterrows():
            job = str(row[requisition_col]).strip()
            skills_text = str(row.get(skill_col, '')).lower()
            skills = [s.strip() for s in re.split(r'[,\n]', skills_text) if s.strip()]
            job_skills[job] = skills

        print(f"SUCCESS: Loaded {len(job_skills)} job requisitions.\n")
    except Exception as e:
        print(f"ERROR: Error loading requisitions: {e}")
        import traceback
        traceback.print_exc()
        return

    resume_files = [f for f in os.listdir(RESUMES_FOLDER) if f.lower().endswith('.docx')]
    print(f"Found {len(resume_files)} resume(s)\n")
    if not resume_files:
        print("WARNING: No .docx files found.")
        return

    all_candidates = []

    for idx, filename in enumerate(resume_files, 1):
        print(f"\n{'='*60}")
        print(f"Processing {idx}/{len(resume_files)}: {filename}")
        print(f"{'='*60}")

        filepath = os.path.join(RESUMES_FOLDER, filename)
        raw_text = extract_docx_text_deep(filepath)
        if not raw_text:
            print("   WARNING: No text extracted, skipping.\n")
            continue

        clean_text = re.sub(r'[\s\xa0]+', ' ', raw_text).strip()
        lower_text = clean_text.lower()

        name = extract_name(clean_text, filename)
        email = extract_email(lower_text)
        phone = extract_phone_robust(clean_text)
        candidate_exp = extract_experience(clean_text)
        skills_section = extract_skills_section(clean_text)

        print(f"\nCandidate: {name}")
        print(f"   Email: {email}")
        print(f"   Phone: {phone}")
        print(f"   Experience: {candidate_exp} years")
        print(f"   Skills section (preview): {skills_section[:200]}...")

        best_job = "No Match"
        best_score = 0.0
        best_skills_matched = []
        best_skill_score = 0.0
        best_exp_score = 0.0
        job_scores = []

        for job, skills in job_skills.items():
            if not skills:
                continue
            print(f"\n   Matching against: {job}")
            matched, skill_score = hybrid_skill_match(
                resume_text=lower_text,
                required_skills=skills,
                skills_section=skills_section.lower(),
                debug=True
            )

            job_exp_min = 0
            if experience_col:
                try:
                    req_row = df_req[df_req[requisition_col] == job]
                    if not req_row.empty and experience_col in req_row.columns:
                        exp_str = str(req_row[experience_col].values[0])
                        job_exp_min = parse_years(exp_str)
                        print(f"   Required experience: {job_exp_min} years")
                except Exception as e:
                    print(f"   Warning: Could not parse experience: {e}")

            exp_score = experience_score(candidate_exp, job_exp_min)
            overall = compute_overall_score(skill_score, exp_score)
            job_scores.append((job, overall, len(matched), skill_score, exp_score))

            if overall > best_score:
                best_score = overall
                best_job = job
                best_skills_matched = [s for s, _ in matched]
                best_skill_score = skill_score
                best_exp_score = exp_score

        job_scores.sort(key=lambda x: x[1], reverse=True)
        print(f"\n   Top matches:")
        for i, (job, score, cnt, s_score, e_score) in enumerate(job_scores[:3]):
            print(f"      {i+1}. {job}: Score {score:.3f} (Skills: {cnt}, Skill Score: {s_score:.2f}, Exp Score: {e_score:.2f})")

        is_shortlisted = best_score >= SHORTLIST_THRESHOLD
        status = "Shortlisted" if is_shortlisted else "Non Shortlisted"
        reason = f"Matched {len(best_skills_matched)} skills with score {best_score:.2f}" if is_shortlisted else f"Score too low ({best_score:.2f})"

        all_candidates.append({
            "Candidate Name": name,
            "Email": email,
            "Mobile": phone,
            "Experience": f"{candidate_exp} Years",
            "Matched Job Position": best_job,
            "Matched Skills": ", ".join(best_skills_matched[:5]) if best_skills_matched else "None",
            "Skill Count": len(best_skills_matched),
            "Skill Score": round(best_skill_score, 3),
            "Experience Score": round(best_exp_score, 3),
            "Score": round(best_score, 3),
            "Status": status,
            "Reason": reason
        })

        print(f"\n   Best match: {best_job} (Score: {best_score:.3f})")
        print(f"   Status: {status}")
        print(f"   Reason: {reason}")

    if not all_candidates:
        print("\nERROR: No candidates processed.")
        return

    df_candidates = pd.DataFrame(all_candidates)
    df_candidates['Rank'] = 0

    for job in df_candidates['Matched Job Position'].unique():
        if job == "No Match":
            continue
        mask = df_candidates['Matched Job Position'] == job
        job_df = df_candidates.loc[mask]
        if not job_df.empty:
            sorted_idx = job_df.sort_values('Score', ascending=False).index
            df_candidates.loc[sorted_idx, 'Rank'] = range(1, len(sorted_idx)+1)

    df_short = df_candidates[df_candidates['Status'] == 'Shortlisted'].copy()
    df_non = df_candidates[df_candidates['Status'] == 'Non Shortlisted'].copy()

    if not df_short.empty:
        df_short = df_short.sort_values(['Matched Job Position', 'Rank'])
    if not df_non.empty:
        df_non = df_non.sort_values('Score', ascending=False)

    try:
        df_short.to_excel(SHORTLISTED_FILE, index=False)
        print(f"\nSUCCESS: Shortlisted saved: {SHORTLISTED_FILE} ({len(df_short)} candidates)")
    except Exception as e:
        print(f"ERROR: Error saving shortlisted: {e}")

    try:
        df_non.to_excel(NON_SHORTLISTED_FILE, index=False)
        print(f"SUCCESS: Non-shortlisted saved: {NON_SHORTLISTED_FILE} ({len(df_non)} candidates)")
    except Exception as e:
        print(f"ERROR: Error saving non-shortlisted: {e}")

    print("\n" + "="*60)
    print("PROCESSING SUMMARY")
    print("="*60)
    print(f"Source: {BASE_PATH}")
    print(f"Resumes processed: {len(resume_files)}")
    print(f"Candidates extracted: {len(all_candidates)}")
    print(f"Shortlisted: {len(df_short)}")
    print(f"Non-shortlisted: {len(df_non)}")
    if len(df_short) > 0:
        print("\nTop matches by job:")
        for job in df_short['Matched Job Position'].unique()[:5]:
            job_candidates = df_short[df_short['Matched Job Position'] == job]
            top = job_candidates.iloc[0]
            print(f"   {job}: #{top['Rank']} {top['Candidate Name']} (Score: {top['Score']})")
    print("="*60)

if __name__ == "__main__":
    main()